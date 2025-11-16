"""
Document Processor
Processes PDF, DOC, DOCX, and TXT files to extract text, metadata, and generate embeddings
"""
import logging
import asyncio
from pathlib import Path
from typing import Dict, Optional, List
import re
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processes document files (PDF, DOC, DOCX, TXT) to extract text and metadata"""
    
    def __init__(self):
        """Initialize document processor with required libraries"""
        self._pdf_available = False
        self._docx_available = False
        
        # Try to import PDF libraries
        self._use_pdfplumber = False
        try:
            import PyPDF2
            self._pdf_available = True
            logger.info("PyPDF2 available for PDF processing")
        except ImportError:
            try:
                import pdfplumber
                self._pdf_available = True
                self._use_pdfplumber = True
                logger.info("pdfplumber available for PDF processing")
            except ImportError:
                logger.warning("No PDF library available. Install PyPDF2 or pdfplumber")
        
        # Try to import DOCX library
        try:
            import docx
            self._docx_available = True
            logger.info("python-docx available for DOCX processing")
        except ImportError:
            logger.warning("python-docx not available. Install python-docx for DOCX support")
    
    async def process(self, file_path: Path, mime_type: str) -> Dict:
        """
        Process document file and return extracted text, metadata, and embeddings
        
        Args:
            file_path: Path to document file
            mime_type: MIME type of the file
            
        Returns:
            Dictionary with text, metadata, category, and embeddings
        """
        try:
            if mime_type == "application/pdf":
                return await self._process_pdf(file_path, mime_type)
            elif mime_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                return await self._process_docx(file_path, mime_type)
            elif mime_type == "text/plain":
                return await self._process_txt(file_path, mime_type)
            else:
                raise ValueError(f"Unsupported document type: {mime_type}")
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}", exc_info=True)
            return {
                "text": "",
                "embeddings": None,
                "category": "uncategorized",
                "metadata": {"error": str(e)},
            }
    
    async def _process_pdf(self, file_path: Path, mime_type: str) -> Dict:
        """Process PDF file to extract text and metadata"""
        if not self._pdf_available:
            return {
                "text": "",
                "embeddings": None,
                "category": "uncategorized",
                "metadata": {"error": "PDF library not available"},
            }
        
        loop = asyncio.get_event_loop()
        
        def extract_pdf_content():
            text_content = []
            metadata = {}
            num_pages = 0
            
            try:
                if hasattr(self, '_use_pdfplumber') and self._use_pdfplumber:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        num_pages = len(pdf.pages)
                        for page in pdf.pages:
                            text = page.extract_text()
                            if text:
                                text_content.append(text)
                        
                        # Extract metadata
                        if pdf.metadata:
                            metadata = {
                                "title": pdf.metadata.get("Title", ""),
                                "author": pdf.metadata.get("Author", ""),
                                "subject": pdf.metadata.get("Subject", ""),
                                "creator": pdf.metadata.get("Creator", ""),
                                "producer": pdf.metadata.get("Producer", ""),
                                "creation_date": str(pdf.metadata.get("CreationDate", "")),
                            }
                else:
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        num_pages = len(pdf_reader.pages)
                        
                        for page in pdf_reader.pages:
                            text = page.extract_text()
                            if text:
                                text_content.append(text)
                        
                        # Extract metadata
                        if pdf_reader.metadata:
                            metadata = {
                                "title": pdf_reader.metadata.get("/Title", ""),
                                "author": pdf_reader.metadata.get("/Author", ""),
                                "subject": pdf_reader.metadata.get("/Subject", ""),
                                "creator": pdf_reader.metadata.get("/Creator", ""),
                                "producer": pdf_reader.metadata.get("/Producer", ""),
                            }
            except Exception as e:
                logger.error(f"Error extracting PDF content: {e}")
                raise
            
            full_text = "\n".join(text_content)
            return full_text, metadata, num_pages
        
        text, metadata, num_pages = await loop.run_in_executor(None, extract_pdf_content)
        
        # Analyze text to determine category
        category = self._categorize_text(text, file_path)
        
        # Generate basic statistics
        word_count = len(text.split())
        char_count = len(text)
        
        metadata.update({
            "mime_type": mime_type,
            "type": "pdf",
            "num_pages": num_pages,
            "word_count": word_count,
            "char_count": char_count,
            "file_size": file_path.stat().st_size,
        })
        
        # Generate simple embeddings (word frequency vector as placeholder)
        # In production, use sentence transformers or similar
        embeddings = self._generate_text_embeddings(text)
        
        return {
            "text": text,
            "embeddings": embeddings,
            "category": category,
            "metadata": metadata,
        }
    
    async def _process_docx(self, file_path: Path, mime_type: str) -> Dict:
        """Process DOCX file to extract text and metadata"""
        if not self._docx_available:
            return {
                "text": "",
                "embeddings": None,
                "category": "uncategorized",
                "metadata": {"error": "DOCX library not available"},
            }
        
        loop = asyncio.get_event_loop()
        
        def extract_docx_content():
            try:
                import docx
                doc = docx.Document(file_path)
                
                # Extract text from all paragraphs
                text_content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                full_text = "\n".join(text_content)
                
                # Extract metadata
                core_props = doc.core_properties
                metadata = {
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "keywords": core_props.keywords or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                }
                
                return full_text, metadata
            except Exception as e:
                logger.error(f"Error extracting DOCX content: {e}")
                raise
        
        text, metadata = await loop.run_in_executor(None, extract_docx_content)
        
        # Analyze text to determine category
        category = self._categorize_text(text, file_path)
        
        # Generate basic statistics
        word_count = len(text.split())
        char_count = len(text)
        
        metadata.update({
            "mime_type": mime_type,
            "type": "docx",
            "word_count": word_count,
            "char_count": char_count,
            "file_size": file_path.stat().st_size,
        })
        
        # Generate embeddings
        embeddings = self._generate_text_embeddings(text)
        
        return {
            "text": text,
            "embeddings": embeddings,
            "category": category,
            "metadata": metadata,
        }
    
    async def _process_txt(self, file_path: Path, mime_type: str) -> Dict:
        """Process TXT file to extract text and metadata"""
        loop = asyncio.get_event_loop()
        
        def read_txt_content():
            try:
                # Try different encodings
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            return f.read()
                    except UnicodeDecodeError:
                        continue
                # If all fail, read as binary and decode with errors='replace'
                with open(file_path, 'rb') as f:
                    return f.read().decode('utf-8', errors='replace')
            except Exception as e:
                logger.error(f"Error reading TXT file: {e}")
                raise
        
        text = await loop.run_in_executor(None, read_txt_content)
        
        # Analyze text to determine category
        category = self._categorize_text(text, file_path)
        
        # Generate basic statistics
        word_count = len(text.split())
        char_count = len(text)
        line_count = len(text.splitlines())
        
        metadata = {
            "mime_type": mime_type,
            "type": "txt",
            "word_count": word_count,
            "char_count": char_count,
            "line_count": line_count,
            "file_size": file_path.stat().st_size,
        }
        
        # Generate embeddings
        embeddings = self._generate_text_embeddings(text)
        
        return {
            "text": text,
            "embeddings": embeddings,
            "category": category,
            "metadata": metadata,
        }
    
    def _categorize_text(self, text: str, file_path: Optional[Path] = None) -> str:
        """
        Categorize document based on text content and filename
        Uses keyword matching to determine document category
        """
        if not text:
            # Try to categorize from filename if text is empty
            if file_path:
                return self._categorize_from_filename(file_path)
            return "uncategorized"
        
        text_lower = text.lower()
        
        # Define category keywords with weighted scoring
        categories = {
            "technical": {
                "keywords": ["code", "function", "api", "algorithm", "software", "programming", "technical", 
                            "implementation", "python", "javascript", "java", "c++", "database", "server",
                            "framework", "library", "module", "class", "method", "variable", "syntax"],
                "weight": 1.0
            },
            "academic": {
                "keywords": ["research", "study", "thesis", "dissertation", "academic", "university", "scholar",
                            "paper", "publication", "journal", "citation", "reference", "hypothesis", "methodology",
                            "analysis", "conclusion", "abstract", "literature review"],
                "weight": 1.2
            },
            "business": {
                "keywords": ["business", "company", "revenue", "profit", "market", "sales", "customer", "strategy",
                            "marketing", "management", "enterprise", "corporate", "client", "product", "service",
                            "quarterly", "annual", "report", "meeting", "presentation"],
                "weight": 1.0
            },
            "legal": {
                "keywords": ["legal", "law", "contract", "agreement", "terms", "conditions", "liability", "jurisdiction",
                            "attorney", "lawyer", "court", "lawsuit", "litigation", "compliance", "regulation",
                            "statute", "clause", "party", "defendant", "plaintiff"],
                "weight": 1.3
            },
            "medical": {
                "keywords": ["medical", "health", "patient", "diagnosis", "treatment", "disease", "symptom", "clinical",
                            "doctor", "physician", "hospital", "medication", "prescription", "therapy", "surgery",
                            "diagnosis", "condition", "disorder", "syndrome", "examination"],
                "weight": 1.2
            },
            "financial": {
                "keywords": ["financial", "finance", "money", "investment", "bank", "account", "transaction", "budget",
                            "stock", "bond", "portfolio", "asset", "liability", "revenue", "expense", "income",
                            "tax", "audit", "accounting", "balance sheet", "cash flow"],
                "weight": 1.1
            },
            "educational": {
                "keywords": ["education", "learning", "course", "lesson", "student", "teacher", "curriculum", "syllabus",
                            "assignment", "homework", "exam", "test", "grade", "semester", "lecture", "tutorial",
                            "textbook", "chapter", "exercise", "quiz"],
                "weight": 1.0
            },
            "literature": {
                "keywords": ["novel", "story", "chapter", "character", "plot", "narrative", "fiction", "literature",
                            "author", "protagonist", "antagonist", "setting", "theme", "dialogue", "prose",
                            "poetry", "poem", "verse", "stanza"],
                "weight": 1.0
            },
            "scientific": {
                "keywords": ["science", "experiment", "hypothesis", "theory", "data", "analysis", "result", "conclusion",
                            "laboratory", "research", "observation", "measurement", "variable", "control", "method",
                            "physics", "chemistry", "biology", "mathematics"],
                "weight": 1.1
            },
            "news": {
                "keywords": ["news", "article", "report", "journalism", "reporter", "headline", "breaking", "update",
                            "event", "incident", "announcement", "press", "media", "coverage"],
                "weight": 1.0
            },
        }
        
        # Score each category
        category_scores = {}
        for category, data in categories.items():
            keywords = data["keywords"]
            weight = data["weight"]
            score = sum(weight for keyword in keywords if keyword in text_lower)
            if score > 0:
                category_scores[category] = score
        
        # Return category with highest score, or try filename if no match
        if category_scores:
            best_category = max(category_scores, key=category_scores.get)
            # Only return if score is above threshold
            if category_scores[best_category] >= 1.0:
                return best_category
        
        # Fallback to filename categorization
        if file_path:
            return self._categorize_from_filename(file_path)
        return "uncategorized"
    
    def _categorize_from_filename(self, file_path: Path) -> str:
        """Categorize document based on filename when text analysis fails"""
        filename_lower = file_path.stem.lower()
        
        filename_keywords = {
            "technical": ["code", "script", "program", "config", "readme", "setup", "install", "tech", "dev"],
            "academic": ["thesis", "dissertation", "paper", "research", "study", "essay", "assignment"],
            "business": ["report", "proposal", "plan", "strategy", "meeting", "presentation", "invoice"],
            "legal": ["contract", "agreement", "terms", "legal", "law", "compliance"],
            "medical": ["medical", "health", "patient", "diagnosis", "prescription"],
            "financial": ["financial", "budget", "invoice", "receipt", "statement", "tax"],
            "educational": ["course", "lesson", "syllabus", "curriculum", "notes", "homework"],
            "literature": ["novel", "story", "book", "chapter", "poem", "poetry"],
            "scientific": ["experiment", "lab", "data", "analysis", "research"],
            "news": ["news", "article", "report", "update"],
        }
        
        for category, keywords in filename_keywords.items():
            for keyword in keywords:
                if keyword in filename_lower:
                    return category
        
        return "document"
    
    def _generate_text_embeddings(self, text: str) -> Optional[List[float]]:
        """
        Generate simple text embeddings
        Uses TF-IDF-like approach as a placeholder
        In production, use sentence-transformers or similar
        """
        if not text:
            return None
        
        # Simple word frequency-based embedding (128 dimensions)
        words = re.findall(r'\b\w+\b', text.lower())
        if not words:
            return None
        
        # Create a simple frequency vector
        # In production, replace with proper sentence embeddings
        word_freq = {}
        for word in words:
            word_freq[word] = word_freq.get(word, 0) + 1
        
        # Normalize and create fixed-size vector (128 dims)
        total_words = len(words)
        embedding = [0.0] * 128
        
        # Use hash-based feature extraction
        for word, freq in word_freq.items():
            # Hash word to get index in embedding vector
            idx = hash(word) % 128
            embedding[idx] += freq / total_words
        
        # Normalize vector
        norm = sum(x * x for x in embedding) ** 0.5
        if norm > 0:
            embedding = [x / norm for x in embedding]
        
        return embedding

